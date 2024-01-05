import deepl
from docx import Document

auth_key = '74ac2384-9b64-3f72-1d50-7dfb83df3f97:fx'

# 字符串文本使用
# text_origin = 'Hi, you know, this work presents  a modeling  and experimental  analysis  on the flame length of buoyant  turbulent  slot  diffusion  flame!'
# translator = deepl.Translator(auth_key)
# text_translated = translator.translate_text(text_origin, target_lang="ZH") # target_lang="EN-US" or "ZH" or "JA" or "FR" or "ES" or "IT" or "NL" or "PL" or "PT-PT" or "PT-BR" or "RU"
# print(f"text_origin: {text_origin}\n"
#       f"text_translated: {text_translated}")
# # translated_text = result.text


# 函数：翻译普通文本
def translate_text(text, target_lang="ZH"):
    translator = deepl.Translator(auth_key)
    # return translator.translate_text(text, target_lang=target_lang).text

    #check the usage of api
    usage = translator.get_usage()
    if usage.any_limit_reached:
        print('Translation limit reached.')
        return f"{translator.translate_text(text, target_lang=target_lang).text}\ntranslation limit reached"
    if usage.character.valid:
        print(
            f"Character usage: {usage.character.count} of {usage.character.limit}")
        return f"{translator.translate_text(text, target_lang=target_lang).text}\n{usage.character.count} of {usage.character.limit}"
    if usage.document.valid:
        print(f"Document usage: {usage.document.count} of {usage.document.limit}")
        return f"{translator.translate_text(text, target_lang=target_lang).text}\n{usage.document.count} of {usage.document.limit}"




# 函数：翻译文本文件
def translate_text_file(file_path, target_lang="ZH", output_file="translated_text.txt"):
    with open(file_path, 'r', encoding='utf-8') as file:
        text = file.read()
    translated_text = translate_text(text, target_lang)
    with open(output_file, 'w', encoding='utf-8') as file:
        file.write(translated_text)

# 函数：翻译 Word 文件
def translate_word_file(file_path, target_lang="ZH", output_file="translated_document.docx"):
    doc = Document(file_path)
    translated_doc = Document()
    for para in doc.paragraphs:
        translated_text = translate_text(para.text, target_lang)
        translated_doc.add_paragraph(translated_text)
    translated_doc.save(output_file)

# 示例使用1
text_origin = 'Hi, you know, this work presents a modeling and experimental analysis on the flame length of buoyant turbulent slot diffusion flame!'
text_translated = translate_text(text_origin, "ZH")
print(f"text_origin: {text_origin}\n"
      f"text_translated: {text_translated}")

# 示例使用2
# translate_text_file("test.txt", "ZH", "translated_text.txt")

# 示例使用3
# translate_word_file("20231030.docx", "ZH", "translated_document-paper-HENG.docx")

