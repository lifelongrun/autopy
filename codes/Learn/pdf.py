import PyPDF2
import re
from docx import Document

def extract_abstract_conclusion(pdf_path):
    """
    Extracts the abstract and conclusion from a scientific paper in PDF format.

    Args:
    pdf_path (str): Path to the PDF file.

    Returns:
    dict: A dictionary containing the abstract and conclusion.
    """
    extracted_info = {"abstract": "", "conclusion": ""}

    # Regular expression patterns for abstract and conclusion
    # Adjust the patterns as necessary based on common structures of your target documents
    abstract_pattern = r"(?s)(?<=\bAbstract\b).*?(?=\bIntroduction\b|\bKeywords\b)"
    conclusion_pattern = r"(?s)(?<=\bConclusions\b).*?(?=\bAcknowledgment\b|\bReferences\b|\bAppendix\b)"

    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)

            # Extracting full text for abstract and conclusion
            full_text = "/n".join([page.extract_text().replace('\n', ' ') for page in reader.pages if page.extract_text()])

            # Extracting abstract and conclusion
            abstract_match = re.search(abstract_pattern, full_text)
            conclusion_match = re.search(conclusion_pattern, full_text)

            if abstract_match:
                extracted_info["abstract"] = abstract_match.group().strip()
            if conclusion_match:
                extracted_info["conclusion"] = conclusion_match.group().strip()

        # Create a new Word document
        doc = Document()
        doc.add_heading('Extracted Information', 0)

        # Add abstract to the document
        doc.add_heading('Abstract:', level=1)
        doc.add_paragraph(extracted_info["abstract"])

        # Add conclusion to the document
        doc.add_heading('Conclusion:', level=1)
        doc.add_paragraph(extracted_info["conclusion"])

        # Save the document
        doc.save(r"extracted_info.docx")

    except Exception as e:
        print(f"Error processing file {pdf_path}: {e}")

    return extracted_info

# Example usage
for pdf_path in ["test2.pdf"]:
    extracted_info = extract_abstract_conclusion(pdf_path)
    for extracted_info in extracted_info.values():
        print(extracted_info)
